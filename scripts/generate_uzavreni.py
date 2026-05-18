import google.generativeai as genai
from docx import Document
from dotenv import load_dotenv
import os
import pathlib

# Načtení API klíče
load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Cesty k souborům
BASE_DIR = pathlib.Path(__file__).parent.parent
PROJEKT_ID = "PROJ_TS_2602_LIS_MAN_ENG"
PROJEKT_DIR = BASE_DIR / "projects" / PROJEKT_ID
PROMPTS_DIR = BASE_DIR / "prompts"
OUTPUT_DIR = PROJEKT_DIR / "dokumenty"
OUTPUT_DIR.mkdir(exist_ok=True)

# Načtení promptu
prompt_text = (PROMPTS_DIR / "uzavreni.md").read_text(encoding="utf-8")

# Načtení kontext.json
kontext = (PROJEKT_DIR / "kontext.json").read_text(encoding="utf-8")

# Sestavení zprávy pro model
user_message = f"""
Vygeneruj Zprávu o ukončení projektu. Dnešní datum je 15.05.2026.

Kontext projektu:
{kontext}
"""

# Příprava vstupních souborů (PDF dokumenty)
pdf_soubory = []
pdf_adresar = PROJEKT_DIR / "podklady"

if pdf_adresar.exists():
    for pdf in pdf_adresar.glob("*.pdf"):
        print(f"Nahrávám: {pdf.name}")
        soubor = genai.upload_file(pdf, mime_type="application/pdf")
        pdf_soubory.append(soubor)

# Volání Gemini API
print("Generuji zprávu o ukončení projektu...")
model = genai.GenerativeModel(
    model_name="gemini-2.5-flash",
    system_instruction=prompt_text
)

obsah = [user_message] + pdf_soubory

response = model.generate_content(
    obsah,
    generation_config=genai.GenerationConfig(
        temperature=0.2
    )
)

vystup = response.text
print("\n--- VÝSTUP ---\n")
print(vystup)

# Uložení do DOCX
doc = Document()
doc.add_heading(f"{PROJEKT_ID} – Zpráva o ukončení projektu", 0)

for radek in vystup.split("\n"):
    if radek.startswith("# "):
        doc.add_heading(radek[2:], 1)
    elif radek.startswith("## "):
        doc.add_heading(radek[3:], 2)
    elif radek.startswith("### "):
        doc.add_heading(radek[4:], 3)
    elif radek.strip():
        doc.add_paragraph(radek)

output_path = OUTPUT_DIR / "Zprava_o_ukonceni_v1.docx"
doc.save(output_path)
print(f"\nDokument uložen: {output_path}")
