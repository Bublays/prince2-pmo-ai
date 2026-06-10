from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
import os
import pathlib
import json
from datetime import date
from copy import deepcopy

# Načtení API klíče
load_dotenv()
client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

# Cesty k souborům
BASE_DIR = pathlib.Path(__file__).parent.parent
PROJEKT_ID = "PROJ_TS_2606_CHMU"  # ← změň pro jiný projekt
PROJEKT_DIR = BASE_DIR / "projects" / PROJEKT_ID
PROMPTS_DIR = BASE_DIR / "prompts"
TEMPLATES_DIR = BASE_DIR / "templates"
OUTPUT_DIR = PROJEKT_DIR / "dokumenty"
OUTPUT_DIR.mkdir(exist_ok=True)

# Načtení kontextu
kontext_path = PROJEKT_DIR / "kontext.json"
kontext_raw = kontext_path.read_text(encoding="utf-8")
kontext = json.loads(kontext_raw)

# Načtení promptu
prompt_text = (PROMPTS_DIR / "charta.md").read_text(encoding="utf-8")

# Datum
dnes = date.today().strftime("%d.%m.%Y")

# Sestavení zprávy pro model — žádáme strukturované sekce
user_message = f"""
Vygeneruj obsah pro Chartu projektu. Dnešní datum je {dnes}.

Kontext projektu:
{kontext_raw}

Výstup musí obsahovat POUZE tyto sekce oddělené značkami.
Každá sekce začíná značkou na samostatném řádku.
Nepiš nic mimo značky.

[MANDАТ]
(2-3 řádky: Jméno – Role – popis odpovědnosti, jeden člověk na řádek)

[CILE_POPIS]
(3-5 vět věcného popisu hlavního cíle projektu)

[RIZIKA]
(seznam rizik, každé na samostatném řádku začínajícím pomlčkou)

[OMEZENI]
(seznam omezení ze strany zadavatele, každé na samostatném řádku začínajícím pomlčkou)

[PREDPOKLADY]
(seznam předpokladů ze strany zadavatele, každé na samostatném řádku začínajícím pomlčkou)

[POPIS_RESENI]
(detailní popis řešení, 5-10 vět)

[PRACOVNI_UKONY]
(seznam plánovaných pracovních úkonů, každý na samostatném řádku začínajícím pomlčkou, chronologicky)

[ROLE_DODAVATEL]
(seznam rolí dodavatele, každá na řádku ve tvaru: Role: Jméno Příjmení)

[ROLE_ZADAVATEL]
(seznam rolí zadavatele, každá na řádku ve tvaru: Role: Jméno Příjmení)
"""

# Volání API
print("Generuji obsah charty...")
response = client.models.generate_content(
    model="gemini-2.5-flash",
    contents=[user_message],
    config=types.GenerateContentConfig(
        system_instruction=prompt_text,
        temperature=0.2
    )
)

vystup = response.text
print("\n--- VÝSTUP AGENTA ---\n")
print(vystup)

# Parsování sekcí
def get_sekce(text, znacka):
    try:
        start = text.index(f"[{znacka}]") + len(f"[{znacka}]")
        try:
            # Najdi další značku
            remaining = text[start:]
            next_tag = None
            import re
            tags = re.finditer(r'\[([A-Z_]+)\]', remaining)
            first_tag = next(tags, None)
            if first_tag:
                konec = start + first_tag.start()
            else:
                konec = len(text)
        except:
            konec = len(text)
        return text[start:konec].strip()
    except ValueError:
        return "k doplnění"

mandat       = get_sekce(vystup, "MANDАT") or get_sekce(vystup, "MANDAT")
cile_popis   = get_sekce(vystup, "CILE_POPIS")
rizika       = get_sekce(vystup, "RIZIKA")
omezeni      = get_sekce(vystup, "OMEZENI")
predpoklady  = get_sekce(vystup, "PREDPOKLADY")
popis_reseni = get_sekce(vystup, "POPIS_RESENI")
ukony        = get_sekce(vystup, "PRACOVNI_UKONY")
role_dod     = get_sekce(vystup, "ROLE_DODAVATEL")
role_zad     = get_sekce(vystup, "ROLE_ZADAVATEL")

# Pomocné funkce pro nahrazování v DOCX
def nahrad_text_v_odstavci(para, stary, novy):
    """Nahradí text v odstavci přes runs, zachová formátování."""
    if stary not in para.text:
        return False
    # Složit celý text odstavce
    full_text = para.text
    if stary not in full_text:
        return False
    # Nahradit v prvním runu, ostatní vymazat
    new_text = full_text.replace(stary, novy)
    for i, run in enumerate(para.runs):
        if i == 0:
            run.text = new_text
        else:
            run.text = ""
    return True

def nahrad_v_doc(doc, stary, novy):
    """Nahradí text ve všech odstavcích a tabulkách."""
    for para in doc.paragraphs:
        nahrad_text_v_odstavci(para, stary, novy)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    nahrad_text_v_odstavci(para, stary, novy)

def nahrad_italic_sekci(doc, klic_slovo, novy_obsah):
    """Najde odstavec obsahující klíčové slovo v kurzívě a nahradí ho."""
    for para in doc.paragraphs:
        if klic_slovo in para.text:
            for run in para.runs:
                run.text = ""
                run.italic = False
            if para.runs:
                para.runs[0].text = novy_obsah
            return True
    return False

def nahrad_bullet_sekci(doc, kotva_text, nova_data):
    """
    Najde odstavec s kotva_text a za nim nahradí bullet body novými.
    nova_data = string s řádky začínajícími '- '
    """
    # Najdi index odstavce s kotvou
    paras = doc.paragraphs
    kotva_idx = None
    for i, para in enumerate(paras):
        if kotva_text in para.text:
            kotva_idx = i
            break
    if kotva_idx is None:
        return False
    # Nahraď existující bullet body za kotvou
    nove_body = [r.lstrip("- ").strip() for r in nova_data.split("\n")
                 if r.strip() and r.strip() != "k doplnění"]
    # Najdi bullet odstavce za kotvou a nahraď je
    bullet_count = 0
    for i in range(kotva_idx + 1, min(kotva_idx + 20, len(paras))):
        para = paras[i]
        if bullet_count < len(nove_body):
            # Nahraď text odstavce
            if para.runs:
                for j, run in enumerate(para.runs):
                    if j == 0:
                        run.text = nove_body[bullet_count]
                        run.italic = False
                    else:
                        run.text = ""
            bullet_count += 1
        else:
            # Vymaž zbývající bullet odstavce
            for run in para.runs:
                run.text = ""
    return True

# Načtení šablony
print("\nVyplňuji šablonu...")
template_path = TEMPLATES_DIR / "charta.docx"
if not template_path.exists():
    # Zkus alternativní název
    for f in TEMPLATES_DIR.glob("*harta*"):
        template_path = f
        break

print(f"Šablona: {template_path}")
doc = Document(template_path)

# Projekt ID
projekt_id = kontext.get("projekt", {}).get("id", PROJEKT_ID)
nahrad_v_doc(doc, "PROJ_XXX_RRMM_TYP", projekt_id)

# Datum
nahrad_v_doc(doc, "17. 1. 2021", dnes)
nahrad_v_doc(doc, "10. 1. 2021", dnes)

# PM a schvalovatel
pm = kontext.get("role", {}).get("projektovy_manazer", "k doplnění")
sponzor = kontext.get("role", {}).get("sponzor_projektu", "k doplnění")
nahrad_v_doc(doc, "Vladislav Poměnka", pm)
nahrad_v_doc(doc, "Martin Voměnka", sponzor)

# Časový rozsah
casovy_rozsah = kontext.get("terminy", {}).get("doba_realizace_dni", "k doplnění")
nejzazsi = kontext.get("terminy", {}).get("nejzazsi_termin", "k doplnění")
nahrad_v_doc(doc, "XY pracovních dnů", f"{casovy_rozsah} dnů")
nahrad_v_doc(doc, "XX. X. 2019", nejzazsi)

# Sekce s kurzívovým textem — nahradit obsahem od agenta
nahrad_italic_sekci(doc, "Stručný popis, obecný", cile_popis)
nahrad_italic_sekci(doc, "detailní popis řešení", popis_reseni)
nahrad_italic_sekci(doc, "printscreen původního harmonogramu", ukony)

# Mandát projektu — nahradit placeholder
nahrad_v_doc(doc, "Jméno příjmení – Pracovní pozice – popis odpovědnosti v projektu (Sponzor projektu)", "")
nahrad_v_doc(doc, "Jméno příjmení – Pracovní pozice – popis odpovědnosti v projektu (Projektový manažer)", "")

# Role dodavatele
nahrad_v_doc(doc, "Jméno Příjmení", "k doplnění")

# Uložení
output_path = OUTPUT_DIR / f"Charta_{PROJEKT_ID}_v1.docx"
doc.save(output_path)
print(f"\nDokument uložen: {output_path}")
print("\nPoznámka: Zkontroluj a doplň sekce označené 'k doplnění'.")
print("Tabulku milníků a role vyplň dle výstupu agenta výše.")
